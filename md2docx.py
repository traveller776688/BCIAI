#!/usr/bin/env python3
"""
Convert FINAL_BLUEPRINT_NLP.md → Word .docx with:
- Rendered math formulas (LaTeX → PNG via matplotlib mathtext)
- Embedded images
- Proper formatting (headings, code, tables, bold)
"""
import re, os, io, hashlib, sys
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

BASE_DIR = r'C:\Users\LENOVO\Desktop\BCIAI'
CACHE_DIR = os.path.join(BASE_DIR, 'math_cache')
os.makedirs(CACHE_DIR, exist_ok=True)

# ── Math rendering ──────────────────────────────────────────
plt.rcParams['mathtext.fontset'] = 'stix'
plt.rcParams['font.size'] = 11

MATH_DPI = 200
MATH_FG = '#1a1a2e'
MATH_BG = '#ffffff'
MATH_PAD = 3  # pixels

def sanitize_latex(latex):
    """Preprocess LaTeX to be compatible with matplotlib's mathtext (TeX subset)."""
    s = latex
    # Remove \big, \Big, \bigg, \Bigg size specifiers (mathtext doesn't support them)
    s = re.sub(r'\\big[lrm]?\s*', '', s)
    s = re.sub(r'\\Big[lrm]?\s*', '', s)
    s = re.sub(r'\\bigg[lrm]?\s*', '', s)
    s = re.sub(r'\\Bigg[lrm]?\s*', '', s)

    # \tfrac{x}{y} → \frac{x}{y} (mathtext supports \frac but not \tfrac)
    s = s.replace(r'\tfrac', r'\frac')

    # \square → Unicode white square (mathtext doesn't support \square or \Box)
    s = s.replace(r'\square', '\u25a1')

    # \operatorname{...} → \mathrm{...} (simpler, supported)
    s = re.sub(r'\\operatorname\{([^}]*)\}', r'\\mathrm{\1}', s)

    # \! \; \, are fine (spacing commands)
    # \| for norm is fine
    # \text{...} is fine
    # \boldsymbol, \mathcal, \mathbb, \mathbf are fine
    # \max, \min, \exp, \log, \ln, \det, \sup, \inf, \lim are fine

    return s


def render_math(latex, display_mode=False):
    """Render LaTeX math to PNG bytes using matplotlib's mathtext."""
    latex_clean = sanitize_latex(latex)

    # Cache key
    key = hashlib.md5(f"{latex_clean}|{display_mode}".encode()).hexdigest()
    cache_path = os.path.join(CACHE_DIR, f'{key}.png')
    if os.path.exists(cache_path):
        with open(cache_path, 'rb') as f:
            return io.BytesIO(f.read())

    try:
        fig, ax = plt.subplots(figsize=(0.01, 0.01), facecolor='none')
        ax.axis('off')

        if display_mode:
            text = ax.text(0.5, 0.5, f'${latex_clean}$', transform=ax.transAxes,
                          fontsize=12, ha='center', va='center',
                          color=MATH_FG)
        else:
            text = ax.text(0, 0, f'${latex_clean}$', fontsize=11,
                          color=MATH_FG)

        fig.canvas.draw()
        bbox = text.get_window_extent(renderer=fig.canvas.get_renderer())
        bbox = bbox.expanded(MATH_PAD, MATH_PAD)

        # Render to image
        fig.set_size_inches(bbox.width/fig.dpi, bbox.height/fig.dpi)
        text.set_position((0.5, 0.5))
        text.set_transform(ax.transAxes)

        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=MATH_DPI, bbox_inches='tight',
                   facecolor=MATH_BG, edgecolor='none', pad_inches=0.05)
        plt.close(fig)
        buf.seek(0)

        # Cache
        with open(cache_path, 'wb') as f:
            f.write(buf.read())
        buf.seek(0)
        return buf

    except Exception as e:
        plt.close('all')
        print(f'  Math render fallback for: {latex[:60]}... ({str(e)[:80]})')
        return None


# ── Document building ───────────────────────────────────────
def build_docx(md_path, output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            for cl in code_lines:
                p = doc.add_paragraph()
                run = p.add_run(cl)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Cm(0.5)
            code_lines = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        ncols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=ncols)
        table.style = 'Light Grid Accent 1'
        for i, row in enumerate(table_rows):
            for j in range(ncols):
                cell_text = row[j].strip() if j < len(row) else ''
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for p_ in cell.paragraphs:
                    for run in p_.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()
        table_rows = []

    def add_image_centered(img_path, caption=''):
        if os.path.exists(img_path):
            try:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_img.add_run()
                run.add_picture(img_path, width=Inches(5.8))
                if caption:
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_clean = caption.strip().strip('*').strip()
                    run_cap = p_cap.add_run(cap_clean)
                    run_cap.font.size = Pt(8.5)
                    run_cap.font.italic = True
            except Exception as e:
                print(f'  WARNING: Could not embed {img_path}: {e}')
        else:
            print(f'  WARNING: Image not found: {img_path}')

    def add_paragraph_with_math(text):
        """Parse text with $...$ math and **bold** and render properly."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

        # Split on display math $$...$$ first
        # Then split on inline $...$
        # Then split on **bold**

        # Strategy: tokenize the line into segments
        segments = tokenize_line(text)
        for seg_type, seg_text in segments:
            if seg_type == 'math_inline':
                buf = render_math(seg_text, display_mode=False)
                if buf:
                    run = p.add_run()
                    run.add_picture(buf, width=Inches(0.22 * max(1, len(seg_text)/4)))
                else:
                    run = p.add_run(f'${seg_text}$')
                    run.font.name = 'Cambria Math'
                    run.font.size = Pt(10)
            elif seg_type == 'math_display':
                # Add as centered paragraph
                buf = render_math(seg_text, display_mode=True)
                if buf:
                    p_disp = doc.add_paragraph()
                    p_disp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_disp.paragraph_format.space_before = Pt(8)
                    p_disp.paragraph_format.space_after = Pt(8)
                    run = p_disp.add_run()
                    img_w = min(5.5, 0.3 * max(1, len(seg_text)/3))
                    run.add_picture(buf, width=Inches(img_w))
                else:
                    p_disp = doc.add_paragraph()
                    p_disp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p_disp.add_run(f'$${seg_text}$$')
                    run.font.name = 'Cambria Math'
            elif seg_type == 'bold':
                run = p.add_run(seg_text)
                run.bold = True
            elif seg_type == 'code':
                run = p.add_run(seg_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            else:
                run = p.add_run(seg_text)

    def tokenize_line(text):
        """Tokenize a line into (type, text) tuples.
        Types: text, math_inline, math_display, bold, code.
        """
        segments = []
        i = 0
        while i < len(text):
            # Display math $$...$$
            if text[i:i+2] == '$$' and i+2 < len(text):
                end = text.find('$$', i+2)
                if end != -1:
                    segments.append(('math_display', text[i+2:end]))
                    i = end + 2
                    continue

            # Inline math $...$
            if text[i] == '$' and i+1 < len(text) and text[i+1] != '$':
                end = text.find('$', i+1)
                if end != -1:
                    segments.append(('math_inline', text[i+1:end]))
                    i = end + 1
                    continue

            # Bold **...**
            if text[i:i+2] == '**' and i+2 < len(text):
                end = text.find('**', i+2)
                if end != -1:
                    segments.append(('bold', text[i+2:end]))
                    i = end + 2
                    continue

            # Inline code `...`
            if text[i] == '`':
                end = text.find('`', i+1)
                if end != -1:
                    segments.append(('code', text[i+1:end]))
                    i = end + 1
                    continue

            # Plain text
            segments.append(('text', text[i]))
            i += 1

        # Merge consecutive text segments
        merged = []
        for seg_type, seg_text in segments:
            if merged and merged[-1][0] == seg_type == 'text':
                merged[-1] = ('text', merged[-1][1] + seg_text)
            else:
                merged.append((seg_type, seg_text))
        return merged


    # ── Main parsing loop ──
    for line in lines:
        stripped = line.rstrip()

        # Image detection
        img_match = re.match(r'^!\[(.*)\]\((.*)\)$', stripped)
        if img_match:
            caption = img_match.group(1)
            rel_path = img_match.group(2)
            img_path = os.path.join(BASE_DIR, rel_path)
            if in_table:
                flush_table()
                in_table = False
            if in_code_block:
                flush_code()
            add_image_centered(img_path, caption)
            print(f'  [IMG] {os.path.basename(rel_path)}')
            continue

        # Code block
        if stripped.startswith('```'):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(stripped)
            continue

        # Empty line
        if not stripped:
            if in_table:
                flush_table()
                in_table = False
            continue

        # Headings
        if stripped.startswith('# '):
            flush_table()
            doc.add_heading(stripped[2:], level=1)
            print(f'  [H1] {stripped[2:60]}')
            continue
        if stripped.startswith('## '):
            flush_table()
            doc.add_heading(stripped[3:], level=2)
            print(f'  [H2] {stripped[3:60]}')
            continue
        if stripped.startswith('### '):
            flush_table()
            doc.add_heading(stripped[4:], level=3)
            continue
        if stripped.startswith('#### '):
            flush_table()
            doc.add_heading(stripped[5:], level=4)
            continue

        # Horizontal rule
        if stripped == '---':
            flush_table()
            doc.add_paragraph('_' * 60)
            continue

        # Table
        if '|' in stripped and stripped.strip().startswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if all(re.match(r'^[-:]+$', c) for c in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue

        # Normal paragraph (with possible math)
        flush_table()
        add_paragraph_with_math(stripped)

    flush_code()
    flush_table()

    doc.save(output_path)
    print(f'\nDone: {output_path}')


if __name__ == '__main__':
    md_path = os.path.join(BASE_DIR, 'FINAL_BLUEPRINT_NLP.md')
    output_path = r'C:\Users\LENOVO\Desktop\FINAL_BLUEPRINT_NLP.docx'
    build_docx(md_path, output_path)
